"""
Quality validation for generated booklet .docx files.

Runs automated checks matching the QC checklist from the production guide,
plus formatting compliance checks from the booklet instructions.
"""

import re
from pathlib import Path
from docx import Document


# Expected sections in a booklet (in rough order)
EXPECTED_SECTIONS = [
    "cover",           # Cover page info
    "recall",          # 20 recall questions
    "knowledge",       # Knowledge chunks (3-5)
    "summary",         # Summary box
    "sentence",        # Sentence starters
    "exam",            # Exam-style questions
    "mark_scheme",     # Mark schemes
    "self_assessment", # Self-assessment grid
]

# Keywords that indicate each section
SECTION_KEYWORDS = {
    "cover": ["self-study booklet", "lesson", "specification"],
    "recall": ["recall", "starter", "one-word", "holistic"],
    "knowledge": ["knowledge chunk", "key vocabulary", "worked example", "misconception", "knowledge check"],
    "summary": ["summary"],
    "sentence": ["sentence starter", "sentence stem", "how to answer", "writing frame"],
    "exam": ["exam-style", "exam style", "exam question", "application question", "calculation question", "application/calculation"],
    "mark_scheme": ["mark scheme", "answer key", "model answer", "grade 4", "grade 7", "grade 9"],
    "self_assessment": ["self-assessment", "self assessment", "score", "total marks"],
}


def validate_docx(docx_path):
    """
    Validate a generated booklet .docx file.

    Returns:
        dict with valid, score, total, checks, warnings, sections_found
    """
    docx_path = Path(docx_path)
    checks = []
    warnings = []

    # ── Check 1: File exists and opens ──
    try:
        doc = Document(str(docx_path))
        checks.append({"name": "file_opens", "passed": True, "detail": "File opens correctly"})
    except Exception as e:
        checks.append({"name": "file_opens", "passed": False, "detail": f"Failed to open: {e}"})
        return {
            "valid": False, "score": 0, "total": 1,
            "checks": checks,
            "warnings": ["File could not be opened — all other checks skipped"],
        }

    # ── Check 2: File size ──
    size_kb = docx_path.stat().st_size / 1024
    size_ok = 5 < size_kb < 10000
    checks.append({
        "name": "file_size", "passed": size_ok,
        "detail": f"File size: {size_kb:.1f} KB (expected 5-10000 KB)",
    })
    if not size_ok:
        warnings.append(f"Unusual file size: {size_kb:.1f} KB")

    # ── Extract all text ──
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()

    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "
    combined_text = full_text_lower + " " + table_text.lower()
    combined_raw = full_text + " " + table_text  # un-lowered for ** check

    # ── Check 3: Word count ──
    word_count = len(full_text.split())
    checks.append({
        "name": "word_count", "passed": word_count > 500,
        "detail": f"Word count: {word_count} (minimum 500)",
    })

    # ── Check 4: Paragraph count ──
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    checks.append({
        "name": "paragraph_count", "passed": para_count > 20,
        "detail": f"Non-empty paragraphs: {para_count} (minimum 20)",
    })

    # ── Check 5: Has tables ──
    table_count = len(doc.tables)
    checks.append({
        "name": "has_tables", "passed": table_count >= 1,
        "detail": f"Tables found: {table_count} (minimum 1)",
    })

    # ── Check 6: Section detection ──
    sections_found = {}
    for section, keywords in SECTION_KEYWORDS.items():
        found = any(kw in combined_text for kw in keywords)
        sections_found[section] = found

    checks.append({
        "name": "has_recall",
        "passed": sections_found.get("recall", False),
        "detail": "Recall starter section " + ("found" if sections_found.get("recall") else "NOT found"),
    })

    chunk_matches = re.findall(r"knowledge\s+chunk|chunk\s+\d", combined_text)
    chunks_ok = len(chunk_matches) >= 2 or sections_found.get("knowledge", False)
    checks.append({
        "name": "has_knowledge_chunks", "passed": chunks_ok,
        "detail": f"Knowledge chunk indicators: {len(chunk_matches)}" +
                  (" (section keywords found)" if sections_found.get("knowledge") else ""),
    })

    checks.append({
        "name": "has_exam_questions",
        "passed": sections_found.get("exam", False),
        "detail": "Exam-style questions section " + ("found" if sections_found.get("exam") else "NOT found"),
    })

    checks.append({
        "name": "has_mark_schemes",
        "passed": sections_found.get("mark_scheme", False),
        "detail": "Mark schemes " + ("found" if sections_found.get("mark_scheme") else "NOT found"),
    })

    checks.append({
        "name": "has_self_assessment",
        "passed": sections_found.get("self_assessment", False),
        "detail": "Self-assessment grid " + ("found" if sections_found.get("self_assessment") else "NOT found"),
    })

    checks.append({
        "name": "has_summary",
        "passed": sections_found.get("summary", False),
        "detail": "Summary section " + ("found" if sections_found.get("summary") else "NOT found"),
    })

    checks.append({
        "name": "has_sentence_starters",
        "passed": sections_found.get("sentence", False),
        "detail": "Sentence starters " + ("found" if sections_found.get("sentence") else "NOT found"),
    })

    # ── Check 7: Headings ──
    heading_count = sum(
        1 for p in doc.paragraphs
        if p.style and p.style.name and p.style.name.startswith("Heading")
    )
    checks.append({
        "name": "has_headings", "passed": heading_count >= 5,
        "detail": f"Headings found: {heading_count} (minimum 5)",
    })

    # ── Check 8: Misconceptions ──
    checks.append({
        "name": "has_misconceptions",
        "passed": "misconception" in combined_text,
        "detail": "Misconception content " + ("found" if "misconception" in combined_text else "NOT found"),
    })

    # ── Check 9: Worked examples ──
    worked_found = "worked example" in combined_text or "example" in combined_text
    checks.append({
        "name": "has_worked_examples", "passed": worked_found,
        "detail": "Worked examples " + ("found" if worked_found else "NOT found"),
    })

    # ══════════════════════════════════════════════════════════════
    # NEW FORMATTING CHECKS
    # ══════════════════════════════════════════════════════════════

    # ── Check 10: No double asterisks ──
    asterisk_count = combined_raw.count("**")
    no_asterisks = asterisk_count == 0
    checks.append({
        "name": "no_double_asterisks", "passed": no_asterisks,
        "detail": f"Double asterisk occurrences: {asterisk_count} (must be 0)",
    })
    if not no_asterisks:
        warnings.append(f"Found {asterisk_count} double-asterisk occurrences — should be 0")

    # ── Check 11: Knowledge content uses bullet style ──
    kc_bullet_ok = _check_section_uses_bullets(doc, "knowledge content")
    checks.append({
        "name": "knowledge_content_bullets", "passed": kc_bullet_ok,
        "detail": "Knowledge Content uses bullet points " +
                  ("correctly" if kc_bullet_ok else "— found numbered list items (should be bullets)"),
    })

    # ── Check 12: Worked examples use bullet style ──
    we_bullet_ok = _check_section_uses_bullets(doc, "worked example")
    checks.append({
        "name": "worked_examples_bullets", "passed": we_bullet_ok,
        "detail": "Worked Examples use bullet points " +
                  ("correctly" if we_bullet_ok else "— found numbered list items (should be bullets)"),
    })

    # ── Check 13: Answer spacing after knowledge check questions ──
    answer_spacing_ok = _check_answer_spacing(doc)
    checks.append({
        "name": "answer_spacing", "passed": answer_spacing_ok,
        "detail": "Knowledge Check answer spacing " +
                  ("adequate" if answer_spacing_ok else "— insufficient blank space between questions"),
    })

    # ── Check 14: PDF companion file exists ──
    pdf_path = docx_path.with_suffix(".pdf")
    pdf_exists = pdf_path.exists()
    checks.append({
        "name": "has_pdf_companion", "passed": pdf_exists,
        "detail": "PDF companion file " + ("found" if pdf_exists else "NOT found"),
    })

    # ── Check 15: Headers/footers present ──
    has_header = False
    try:
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text = "".join(p.text for p in section.header.paragraphs).strip()
                if header_text:
                    has_header = True
                    break
    except Exception:
        pass
    checks.append({
        "name": "has_headers_footers", "passed": has_header,
        "detail": "Document header " + ("found" if has_header else "NOT found"),
    })

    # ── Check 16: UK English spelling ──
    us_spelling_found = _check_us_spellings(combined_text)
    uk_ok = len(us_spelling_found) == 0
    checks.append({
        "name": "uk_english", "passed": uk_ok,
        "detail": "UK English spelling " + (
            "OK" if uk_ok
            else f"— found US spellings: {', '.join(us_spelling_found[:5])}"
        ),
    })
    if not uk_ok:
        warnings.append(f"US English spellings detected: {', '.join(us_spelling_found[:10])}")

    # ── Check 17: Application questions section exists ──
    has_app_q = any(kw in combined_text for kw in [
        "application question", "calculation question",
    ])
    checks.append({
        "name": "has_application_questions", "passed": has_app_q,
        "detail": "Application/Calculation Questions " + ("found" if has_app_q else "NOT found"),
    })

    # ── Check 18: Answer spacing after application questions ──
    app_spacing_ok = _check_answer_spacing_section(doc, "application")
    checks.append({
        "name": "application_answer_spacing", "passed": app_spacing_ok,
        "detail": "Application Questions answer spacing " +
                  ("adequate" if app_spacing_ok else "— insufficient blank space between questions"),
    })

    # ── Summary ──
    passed = sum(1 for c in checks if c["passed"])
    total = len(checks)
    valid = passed >= total * 0.7  # 70% threshold

    return {
        "valid": valid,
        "score": passed,
        "total": total,
        "checks": checks,
        "warnings": warnings,
        "sections_found": sections_found,
    }


def _check_section_uses_bullets(doc, section_keyword):
    """
    Check that paragraphs after a section heading containing `section_keyword`
    use bullet style rather than numbered list style.
    Returns True if compliant (no numbered items found in that section).
    """
    in_section = False
    for p in doc.paragraphs:
        # Detect section heading
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            if section_keyword in p.text.lower():
                in_section = True
                continue
            elif in_section:
                # New heading = left section
                in_section = False

        if in_section:
            style_name = (p.style.name if p.style and p.style.name else "").lower()
            if "list number" in style_name:
                return False

    return True


def _check_answer_spacing(doc):
    """
    Check that after numbered items in a 'Knowledge Check' section,
    there are blank paragraphs for answer writing space.
    Returns True if at least some spacing is found.
    """
    in_kc = False
    found_question = False
    spacing_after_question = 0

    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            if "knowledge check" in p.text.lower():
                in_kc = True
                found_question = False
                continue
            elif in_kc:
                in_kc = False

        if in_kc:
            style_name = (p.style.name if p.style and p.style.name else "").lower()
            if "list number" in style_name:
                if found_question and spacing_after_question < 2:
                    return False  # previous question had no spacing
                found_question = True
                spacing_after_question = 0
            elif found_question and not p.text.strip():
                spacing_after_question += 1

    return True


def _check_answer_spacing_section(doc, section_keyword):
    """
    Check that after numbered items in a section containing `section_keyword`,
    there are blank paragraphs for answer writing space.
    Returns True if at least some spacing is found (or section not found).
    """
    in_section = False
    found_question = False
    spacing_after_question = 0
    section_exists = False

    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            if section_keyword in p.text.lower():
                in_section = True
                section_exists = True
                found_question = False
                continue
            elif in_section:
                in_section = False

        if in_section:
            style_name = (p.style.name if p.style and p.style.name else "").lower()
            if "list number" in style_name:
                if found_question and spacing_after_question < 2:
                    return False
                found_question = True
                spacing_after_question = 0
            elif found_question and not p.text.strip():
                spacing_after_question += 1

    # If section doesn't exist, pass the check (not applicable)
    if not section_exists:
        return True
    return True


# Common US spellings to detect (case-insensitive)
_US_SPELLING_PATTERNS = [
    (r"\borganize\b", "organize→organise"),
    (r"\borganized\b", "organized→organised"),
    (r"\brecognize\b", "recognize→recognise"),
    (r"\brecognized\b", "recognized→recognised"),
    (r"\bminimize\b", "minimize→minimise"),
    (r"\bmaximize\b", "maximize→maximise"),
    (r"\bspecialize\b", "specialize→specialise"),
    (r"\bspecialized\b", "specialized→specialised"),
    (r"\banalyze\b", "analyze→analyse"),
    (r"\banalyzed\b", "analyzed→analysed"),
    (r"\bsummarize\b", "summarize→summarise"),
    (r"\bneutralize\b", "neutralize→neutralise"),
    (r"\boxidize\b", "oxidize→oxidise"),
    (r"\bionize\b", "ionize→ionise"),
    (r"\bcolor\b", "color→colour"),
    (r"\bcolors\b", "colors→colours"),
    (r"\bfavor\b", "favor→favour"),
    (r"\bbehavior\b", "behavior→behaviour"),
    (r"\bcenter\b", "center→centre"),
    (r"\bcenters\b", "centers→centres"),
    (r"\bfiber\b", "fiber→fibre"),
    (r"\bfibers\b", "fibers→fibres"),
    (r"\bliter\b", "liter→litre"),
    (r"\bmeter\b", "meter→metre"),
    (r"\blabeled\b", "labeled→labelled"),
    (r"\bmodeling\b", "modeling→modelling"),
    (r"\bdefense\b", "defense→defence"),
    (r"\bgray\b", "gray→grey"),
    (r"\bsulfur\b", "sulfur→sulphur"),
    (r"\bsulfate\b", "sulfate→sulphate"),
    (r"\baluminum\b", "aluminum→aluminium"),
    (r"\bhemoglobin\b", "hemoglobin→haemoglobin"),
    (r"\bestrogen\b", "estrogen→oestrogen"),
    (r"\bfetus\b", "fetus→foetus"),
    (r"\besophagus\b", "esophagus→oesophagus"),
    (r"\bdiarrhea\b", "diarrhea→diarrhoea"),
    (r"\banemia\b", "anemia→anaemia"),
]


def _check_us_spellings(text):
    """Check for US English spellings in text. Returns list of findings."""
    findings = []
    for pattern, label in _US_SPELLING_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            findings.append(label)
    return findings


def validate_markdown(md_path):
    """Quick validation on markdown content before docx conversion."""
    md_path = Path(md_path)
    content = md_path.read_text()
    content_lower = content.lower()

    checks = []

    word_count = len(content.split())
    checks.append({
        "name": "word_count", "passed": word_count > 1000,
        "detail": f"Word count: {word_count} (minimum 1000 for markdown)",
    })

    for section, keywords in SECTION_KEYWORDS.items():
        found = any(kw in content_lower for kw in keywords)
        checks.append({
            "name": f"section_{section}", "passed": found,
            "detail": f"{section}: {'found' if found else 'NOT found'}",
        })

    table_lines = [l for l in content.split("\n") if l.strip().startswith("|")]
    checks.append({
        "name": "has_tables", "passed": len(table_lines) > 5,
        "detail": f"Table lines: {len(table_lines)}",
    })

    # Check for double asterisks in markdown
    asterisk_count = len(re.findall(r"\*\*", content))
    checks.append({
        "name": "no_double_asterisks", "passed": asterisk_count == 0,
        "detail": f"Double asterisks in markdown: {asterisk_count} (must be 0)",
    })

    passed = sum(1 for c in checks if c["passed"])
    total = len(checks)

    return {
        "valid": passed >= total * 0.7,
        "score": passed,
        "total": total,
        "checks": checks,
    }
