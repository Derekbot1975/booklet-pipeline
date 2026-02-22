"""
Scheme of work export: generate Word documents from parsed lesson data.
"""
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def export_scheme_docx(config, lessons):
    """
    Generate a formatted .docx containing the full scheme of work.

    Args:
        config: course config dict
        lessons: list of all lesson dicts (from get_data()["all_lessons"])

    Returns:
        str path to the generated .docx file
    """
    doc = Document()

    # Page setup — landscape A4
    for section in doc.sections:
        section.orientation = 1  # Landscape
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{config['name']} — Scheme of Work")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata line
    meta_parts = []
    eb = config.get("exam_board")
    if eb and eb != "None":
        meta_parts.append(f"Exam Board: {eb}")
    qual = config.get("qualification")
    if qual and qual != "None":
        meta_parts.append(f"Qualification: {qual}")
    ks = config.get("key_stage")
    if ks:
        meta_parts.append(f"Key Stage: {ks}")
    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        meta_para.runs[0].font.size = Pt(11)
        meta_para.runs[0].font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

    # Group lessons by year
    years = {}
    for lesson in lessons:
        yr = lesson["year"]
        years.setdefault(yr, []).append(lesson)

    # Determine which optional columns have data
    col_map = config.get("col_map", {})
    has_rp = "rp" in col_map
    has_ht = "ht_only" in col_map

    # Build column definitions: (header_text, field_key, width_cm)
    columns = [
        ("Wk", "week", 1.2),
        ("L#", "lesson_number", 1.2),
        ("Subject", "subject", 2.5),
        ("Topic", "topic", 3.5),
        ("Title", "title", 5.0),
        ("Specification Content", "spec_content", 8.0),
        ("Key Vocabulary", "key_vocabulary", 4.0),
    ]
    if has_rp:
        columns.append(("RP", "required_practical", 1.5))
    if has_ht:
        columns.append(("HT", "ht_only", 1.5))

    for yr in sorted(years.keys()):
        # Year heading
        yr_para = doc.add_paragraph()
        yr_run = yr_para.add_run(f"Year {yr}")
        yr_run.bold = True
        yr_run.font.size = Pt(16)
        yr_run.font.color.rgb = RGBColor(0x22, 0x66, 0xAA)

        yr_lessons = years[yr]

        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Light Grid Accent 1"
        table.autofit = True

        # Set column widths
        for i, (_, _, width) in enumerate(columns):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Header row
        hdr_row = table.rows[0]
        for i, (header_text, _, _) in enumerate(columns):
            cell = hdr_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Dark blue background
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "2266AA",
            })
            shading.append(shd)

        # Data rows
        for lesson in yr_lessons:
            row = table.add_row()
            for i, (_, field, _) in enumerate(columns):
                val = lesson.get(field)
                cell_text = str(val) if val is not None else ""
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.size = Pt(7)

            # Highlight non-booklet lessons in light grey
            if not lesson.get("is_booklet_lesson", True):
                for i in range(len(columns)):
                    shading = row.cells[i]._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "F0F0F0",
                    })
                    shading.append(shd)

        # Add some space after the table
        doc.add_paragraph()

    # Summary at the end
    total = len(lessons)
    booklet_count = sum(1 for l in lessons if l.get("is_booklet_lesson", False))
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(
        f"Total lessons: {total} | Booklet lessons: {booklet_count} | "
        f"Non-booklet: {total - booklet_count}"
    )
    summary_run.font.size = Pt(9)
    summary_run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name
