"""
Assessment Builder (Prompt Sheet 17).

Generates end-of-unit tests and custom assessments with:
  - Multiple choice questions (1 mark each)
  - Short answer questions (2-3 marks each)
  - Long answer / extended writing questions (4-6 marks each)
  - Full mark schemes
  - Question bank management

Questions are stored in a JSON-based question bank for reuse.
"""

import json
import logging
import os
import re
import time
import uuid
from datetime import datetime
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv(override=True)

logger = logging.getLogger(__name__)


def _extract_json(raw, expect_object=False):
    """Robustly extract JSON from AI response text."""
    # Strip markdown fences
    raw = re.sub(r"^```(?:json)?\s*\n?", "", raw)
    raw = re.sub(r"\n?```\s*$", "", raw)
    # Find the outermost JSON structure
    if expect_object:
        start = raw.find("{")
        end = raw.rfind("}")
    else:
        start = raw.find("[")
        end = raw.rfind("]")
    if start != -1 and end != -1 and end > start:
        raw = raw[start:end + 1]
    # Fix trailing commas
    raw = re.sub(r",\s*([}\]])", r"\1", raw)
    return raw


ASSESSMENTS_DIR = Path(__file__).parent / "data" / "assessments"
ASSESSMENTS_DIR.mkdir(parents=True, exist_ok=True)

QUESTION_BANK_DIR = Path(__file__).parent / "data" / "question-bank"
QUESTION_BANK_DIR.mkdir(parents=True, exist_ok=True)


ASSESSMENT_SYSTEM_PROMPT = """You are an expert UK assessment writer creating tests for students.
Generate questions that genuinely test understanding, not just recall.

LANGUAGE: Use UK English throughout.

OUTPUT: You MUST return valid JSON matching the schema in the user prompt.
Do NOT wrap the JSON in markdown code fences. Return raw JSON only.

QUESTION QUALITY RULES:
- Multiple choice distractors must reflect common misconceptions
- Short answer questions test understanding, not definitions
- Long answer questions require application, analysis, or evaluation
- Use appropriate command words: Explain, Compare, Evaluate, Describe, Analyse
- Mark schemes must be clear with accept alternatives noted
- All content must be accurate and curriculum-appropriate
- Questions should cover the breadth of the selected content
- Include a mix of recall, understanding, and application"""


def _get_client():
    return anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))


# ───────────────────────────────────────────────────────────────────
# Assessment generation
# ───────────────────────────────────────────────────────────────────

def generate_assessment(lessons, subject, year_group, config,
                        num_mc=10, num_short=8, num_long=3,
                        difficulty="medium", assessment_type="custom",
                        reference_context="",
                        model="claude-sonnet-4-5-20250929"):
    """
    Generate a complete assessment from selected lessons.

    Args:
        lessons: List of lesson dicts to assess
        subject: Subject name
        year_group: Year group number
        config: Dict with options (include_mark_scheme, include_answer_spaces, etc.)
        num_mc: Number of multiple choice questions
        num_short: Number of short answer questions
        num_long: Number of long answer questions
        difficulty: "easy" | "medium" | "hard"
        assessment_type: "end_of_unit" | "custom" | "mock_exam"
        reference_context: Reference document context
        model: Claude model to use

    Returns:
        Assessment data dict
    """
    client = _get_client()
    start = time.time()

    # Build lesson content summary
    lesson_content = []
    for l in lessons:
        lesson_content.append(
            f"L{l.get('lesson_number', '?')}: {l.get('title', 'Untitled')}\n"
            f"  Spec: {l.get('spec_content', 'N/A')}\n"
            f"  Vocab: {l.get('key_vocabulary', 'N/A')}"
        )

    total_mc_marks = num_mc * 1
    total_short_marks = num_short * 2  # average 2 marks
    total_long_marks = num_long * 5    # average 5 marks
    total_marks = total_mc_marks + total_short_marks + total_long_marks
    est_time = max(30, (total_marks * 1.2))  # roughly 1.2 min per mark

    user_prompt = f"""Generate a complete assessment for Year {year_group} {subject}.

ASSESSMENT TYPE: {assessment_type.replace("_", " ").title()}
DIFFICULTY: {difficulty}

CONTENT TO ASSESS:
{chr(10).join(lesson_content)}

{f"REFERENCE MATERIALS:{chr(10)}{reference_context}" if reference_context else ""}

QUESTION REQUIREMENTS:
- Multiple Choice: {num_mc} questions (1 mark each)
- Short Answer: {num_short} questions (2-3 marks each)
- Long Answer: {num_long} questions (4-6 marks each)

TOTAL MARKS: approximately {total_marks}
ESTIMATED TIME: {int(est_time)} minutes

Return the assessment as JSON:
{{
    "title": "Year {year_group} {subject} — [Assessment Name]",
    "subject": "{subject}",
    "yearGroup": {year_group},
    "totalMarks": <actual total>,
    "estimatedTime": "{int(est_time)} minutes",
    "assessmentType": "{assessment_type}",
    "sections": [
        {{
            "name": "Section A: Multiple Choice",
            "instructions": "Choose the correct answer for each question.",
            "totalMarks": {total_mc_marks},
            "questions": [
                {{
                    "number": 1,
                    "questionText": "Which organelle...",
                    "type": "multiple_choice",
                    "options": ["A) Cell wall", "B) Nucleus", "C) Mitochondria", "D) Cell membrane"],
                    "marks": 1,
                    "topic": "Cell Structure",
                    "difficulty": "easy",
                    "markScheme": {{
                        "correctAnswer": "B",
                        "explanation": "The nucleus contains DNA and controls cell activities."
                    }}
                }}
            ]
        }},
        {{
            "name": "Section B: Short Answer",
            "instructions": "Answer each question in the space provided.",
            "totalMarks": <actual total>,
            "questions": [
                {{
                    "number": {num_mc + 1},
                    "questionText": "Explain why...",
                    "type": "short_answer",
                    "marks": 2,
                    "topic": "Topic name",
                    "difficulty": "medium",
                    "markScheme": {{
                        "points": ["Point 1 (1 mark)", "Point 2 (1 mark)"],
                        "acceptAlternatives": "Accept equivalent wording"
                    }}
                }}
            ]
        }},
        {{
            "name": "Section C: Extended Writing",
            "instructions": "Write your answers in full sentences. Use scientific terminology.",
            "totalMarks": <actual total>,
            "questions": [
                {{
                    "number": {num_mc + num_short + 1},
                    "questionText": "Compare and contrast...",
                    "type": "long_answer",
                    "marks": 6,
                    "topic": "Topic name",
                    "difficulty": "hard",
                    "markScheme": {{
                        "indicativeContent": ["Key point 1", "Key point 2"],
                        "levels": [
                            {{"level": 3, "marks": "5-6", "descriptor": "Detailed comparison with accurate use of terminology"}},
                            {{"level": 2, "marks": "3-4", "descriptor": "Some comparison with some terminology"}},
                            {{"level": 1, "marks": "1-2", "descriptor": "Basic points with limited terminology"}}
                        ]
                    }}
                }}
            ]
        }}
    ],
    "gradeBoundaries": {{
        "distinction": <int>,
        "merit": <int>,
        "pass": <int>
    }}
}}"""

    from ai_client import create_message
    response = create_message(
        model=model,
        max_tokens=10000,
        system=[{
            "type": "text",
            "text": ASSESSMENT_SYSTEM_PROMPT,
            "cache_control": {"type": "ephemeral"},
        }],
        messages=[{"role": "user", "content": user_prompt}],
    )

    duration = round(time.time() - start, 1)
    raw = response.content[0].text.strip()

    raw = _extract_json(raw, expect_object=True)

    try:
        assessment = json.loads(raw)
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse assessment JSON: {e}\nRaw: {raw[:800]}")
        raise ValueError(f"AI generated invalid JSON for assessment: {e}")

    # Add metadata
    assessment["id"] = str(uuid.uuid4())[:8]
    assessment["created_at"] = datetime.utcnow().isoformat()
    assessment["generation"] = {
        "model": model,
        "duration_s": duration,
        "usage": {
            "input_tokens": response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        },
    }
    assessment["lesson_refs"] = [
        {"year": l.get("year"), "lesson_number": l.get("lesson_number"), "title": l.get("title")}
        for l in lessons
    ]

    # Save assessment
    path = ASSESSMENTS_DIR / f"{assessment['id']}.json"
    path.write_text(json.dumps(assessment, indent=2))

    # Add questions to the question bank
    _add_to_question_bank(assessment, subject)

    return assessment


# ───────────────────────────────────────────────────────────────────
# Assessment CRUD
# ───────────────────────────────────────────────────────────────────

def list_assessments(subject=None):
    """List all saved assessments."""
    assessments = []
    for p in sorted(ASSESSMENTS_DIR.glob("*.json")):
        try:
            data = json.loads(p.read_text())
            if subject and data.get("subject", "").lower() != subject.lower():
                continue
            assessments.append({
                "id": data["id"],
                "title": data.get("title", ""),
                "subject": data.get("subject", ""),
                "yearGroup": data.get("yearGroup", ""),
                "totalMarks": data.get("totalMarks", 0),
                "estimatedTime": data.get("estimatedTime", ""),
                "assessmentType": data.get("assessmentType", ""),
                "created_at": data.get("created_at", ""),
                "question_count": sum(
                    len(s.get("questions", []))
                    for s in data.get("sections", [])
                ),
            })
        except (json.JSONDecodeError, KeyError):
            continue
    return assessments


def get_assessment(assessment_id):
    """Get a full assessment by ID."""
    path = ASSESSMENTS_DIR / f"{assessment_id}.json"
    if not path.exists():
        return None
    return json.loads(path.read_text())


def update_assessment(assessment_id, updates):
    """Update an assessment (e.g. edit a question)."""
    path = ASSESSMENTS_DIR / f"{assessment_id}.json"
    if not path.exists():
        raise ValueError(f"Assessment not found: {assessment_id}")

    data = json.loads(path.read_text())
    data.update(updates)
    data["updated_at"] = datetime.utcnow().isoformat()
    path.write_text(json.dumps(data, indent=2))
    return data


def delete_assessment(assessment_id):
    """Delete an assessment."""
    path = ASSESSMENTS_DIR / f"{assessment_id}.json"
    if path.exists():
        path.unlink()
        return True
    return False


# ───────────────────────────────────────────────────────────────────
# Question Bank
# ───────────────────────────────────────────────────────────────────

def _add_to_question_bank(assessment, subject):
    """Add all questions from an assessment to the question bank."""
    for section in assessment.get("sections", []):
        for q in section.get("questions", []):
            q_id = str(uuid.uuid4())[:8]
            bank_entry = {
                "id": q_id,
                "subject": subject,
                "topic": q.get("topic", ""),
                "question_type": q.get("type", ""),
                "question_text": q.get("questionText", ""),
                "options": q.get("options"),
                "marks": q.get("marks", 1),
                "difficulty": q.get("difficulty", "medium"),
                "mark_scheme": q.get("markScheme", {}),
                "times_used": 1,
                "is_starred": False,
                "source_assessment_id": assessment.get("id", ""),
                "created_at": datetime.utcnow().isoformat(),
            }
            path = QUESTION_BANK_DIR / f"{q_id}.json"
            path.write_text(json.dumps(bank_entry, indent=2))


def list_question_bank(subject=None, topic=None, question_type=None,
                       difficulty=None, starred_only=False):
    """Browse the question bank with filters."""
    questions = []
    for p in sorted(QUESTION_BANK_DIR.glob("*.json")):
        try:
            q = json.loads(p.read_text())
        except (json.JSONDecodeError, KeyError):
            continue

        if subject and q.get("subject", "").lower() != subject.lower():
            continue
        if topic and topic.lower() not in q.get("topic", "").lower():
            continue
        if question_type and q.get("question_type") != question_type:
            continue
        if difficulty and q.get("difficulty") != difficulty:
            continue
        if starred_only and not q.get("is_starred"):
            continue

        questions.append(q)

    return questions


def star_question(question_id, starred=True):
    """Star or unstar a question in the bank."""
    path = QUESTION_BANK_DIR / f"{question_id}.json"
    if not path.exists():
        raise ValueError(f"Question not found: {question_id}")
    q = json.loads(path.read_text())
    q["is_starred"] = starred
    path.write_text(json.dumps(q, indent=2))
    return q


def get_question(question_id):
    """Get a single question from the bank."""
    path = QUESTION_BANK_DIR / f"{question_id}.json"
    if not path.exists():
        return None
    return json.loads(path.read_text())


# ───────────────────────────────────────────────────────────────────
# Export as DOCX
# ───────────────────────────────────────────────────────────────────

def export_student_paper(assessment_id, output_path=None):
    """Export the student paper (questions only, no mark scheme) as DOCX."""
    assessment = get_assessment(assessment_id)
    if not assessment:
        raise ValueError(f"Assessment not found: {assessment_id}")

    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading(assessment.get("title", "Assessment"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Total Marks: {assessment.get('totalMarks', '')}    "
        f"Time: {assessment.get('estimatedTime', '')}    "
        f"Date: _______________"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Name field
    name_para = doc.add_paragraph()
    run = name_para.add_run("Name: _________________________________    Class: __________")
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    for section in assessment.get("sections", []):
        # Section heading
        h = doc.add_heading(section.get("name", ""), level=2)

        # Instructions
        inst = doc.add_paragraph()
        run = inst.add_run(section.get("instructions", ""))
        run.font.size = Pt(11)
        run.italic = True

        for q in section.get("questions", []):
            q_num = q.get("number", "")
            q_text = q.get("questionText", "")
            marks = q.get("marks", 1)
            q_type = q.get("type", "")

            # Question text
            p = doc.add_paragraph()
            run = p.add_run(f"{q_num}. {q_text}")
            run.font.size = Pt(12)
            run.bold = True

            # Mark allocation
            mark_run = p.add_run(f"  [{marks} mark{'s' if marks > 1 else ''}]")
            mark_run.font.size = Pt(10)
            mark_run.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

            if q_type == "multiple_choice" and q.get("options"):
                for opt in q["options"]:
                    op = doc.add_paragraph(f"    {opt}")
                    op.paragraph_format.space_before = Pt(2)
                    op.paragraph_format.space_after = Pt(2)
                    for run in op.runs:
                        run.font.size = Pt(12)
            elif q_type == "short_answer":
                # Answer lines
                for _ in range(max(2, marks)):
                    line = doc.add_paragraph("_" * 80)
                    line.paragraph_format.space_before = Pt(6)
                    for run in line.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            elif q_type == "long_answer":
                # More answer lines
                for _ in range(max(4, marks + 2)):
                    line = doc.add_paragraph("_" * 80)
                    line.paragraph_format.space_before = Pt(6)
                    for run in line.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

            doc.add_paragraph()  # spacing

    if not output_path:
        output_dir = Path(__file__).parent / "output" / "assessments"
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_title = re.sub(r"[^\w\s-]", "", assessment.get("title", "assessment")).replace(" ", "_")
        output_path = str(output_dir / f"{safe_title} - Student Paper.docx")

    doc.save(output_path)
    return output_path


def export_mark_scheme(assessment_id, output_path=None):
    """Export the mark scheme as DOCX."""
    assessment = get_assessment(assessment_id)
    if not assessment:
        raise ValueError(f"Assessment not found: {assessment_id}")

    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading(f"{assessment.get('title', 'Assessment')} — Mark Scheme", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Total Marks: {assessment.get('totalMarks', '')}")
    run.font.size = Pt(11)

    doc.add_paragraph()

    for section_data in assessment.get("sections", []):
        doc.add_heading(section_data.get("name", ""), level=2)

        for q in section_data.get("questions", []):
            q_num = q.get("number", "")
            q_text = q.get("questionText", "")
            marks = q.get("marks", 1)
            ms = q.get("markScheme", {})

            # Question
            p = doc.add_paragraph()
            run = p.add_run(f"Q{q_num}. {q_text} [{marks} marks]")
            run.font.size = Pt(11)
            run.bold = True

            # Mark scheme content
            if q.get("type") == "multiple_choice":
                ans = doc.add_paragraph()
                run = ans.add_run(f"Answer: {ms.get('correctAnswer', '')}")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                if ms.get("explanation"):
                    exp = doc.add_paragraph()
                    run = exp.add_run(f"Explanation: {ms['explanation']}")
                    run.font.size = Pt(10)
                    run.italic = True

            elif q.get("type") == "short_answer":
                for point in ms.get("points", []):
                    pp = doc.add_paragraph(f"• {point}")
                    for run in pp.runs:
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                if ms.get("acceptAlternatives"):
                    alt = doc.add_paragraph()
                    run = alt.add_run(f"Accept: {ms['acceptAlternatives']}")
                    run.font.size = Pt(10)
                    run.italic = True

            elif q.get("type") == "long_answer":
                if ms.get("indicativeContent"):
                    doc.add_paragraph("Indicative content:")
                    for ic in ms["indicativeContent"]:
                        pp = doc.add_paragraph(f"• {ic}")
                        for run in pp.runs:
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                if ms.get("levels"):
                    doc.add_paragraph("Levels of response:")
                    for lvl in ms["levels"]:
                        pp = doc.add_paragraph(
                            f"Level {lvl.get('level', '?')} ({lvl.get('marks', '')} marks): "
                            f"{lvl.get('descriptor', '')}"
                        )
                        for run in pp.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph()

    # Grade boundaries
    gb = assessment.get("gradeBoundaries", {})
    if gb:
        doc.add_heading("Grade Boundaries", level=2)
        for grade, mark in sorted(gb.items(), key=lambda x: -x[1] if isinstance(x[1], (int, float)) else 0):
            p = doc.add_paragraph(f"{grade.title()}: {mark} marks")
            for run in p.runs:
                run.font.size = Pt(11)

    if not output_path:
        output_dir = Path(__file__).parent / "output" / "assessments"
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_title = re.sub(r"[^\w\s-]", "", assessment.get("title", "assessment")).replace(" ", "_")
        output_path = str(output_dir / f"{safe_title} - Mark Scheme.docx")

    doc.save(output_path)
    return output_path
