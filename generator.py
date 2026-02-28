"""
Claude API integration for automated booklet generation.

Sends the master prompt to Claude and receives the generated booklet content.
Includes markdown sanitisation, .docx conversion with full formatting,
and PDF export.
"""

import logging
import os
import re
import time
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from dotenv import load_dotenv

import anthropic

load_dotenv(override=True)

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).parent / "output"

# ---------------------------------------------------------------------------
# Formatting constants  (Booklet Formatting Specification v1.0)
# ---------------------------------------------------------------------------

FONT = "Arial"

# Colour objects for font runs
COL_PRIMARY   = RGBColor(0x1A, 0x1A, 0x1A)
COL_SECONDARY = RGBColor(0x33, 0x33, 0x33)
COL_MUTED     = RGBColor(0x66, 0x66, 0x66)
COL_MISC_TTL  = RGBColor(0xCC, 0x00, 0x00)
COL_WE_LABEL  = RGBColor(0x22, 0x66, 0xAA)

# Hex strings for borders / cell shading
HEX_BORDER_LT  = "CCCCCC"
HEX_BORDER_MD  = "999999"
HEX_TBL_HDR_BG = "F0F0F0"
HEX_MISC_BDR   = "CC0000"
HEX_MISC_BG    = "FFF5F5"
HEX_WE_BDR     = "2266AA"
HEX_WE_BG      = "F5F8FF"

# (size_pt, colour, space_before_pt, space_after_pt)
HEADING_SPEC = {
    1: (26, COL_PRIMARY,   0,  6),
    2: (18, COL_PRIMARY,  24, 12),
    3: (14, COL_SECONDARY, 18, 8),
    4: (12, COL_SECONDARY, 14, 6),
    5: (11, COL_SECONDARY, 10, 4),
}

TEXT_AREA_MM = 160  # 210 − 25 − 25

# ---------------------------------------------------------------------------
# SYSTEM PROMPT — comprehensive formatting instructions for Claude
# ---------------------------------------------------------------------------

DEFAULT_SYSTEM_PROMPT_CONTEXT = "AQA GCSE Combined Science: Trilogy (8464)"

SYSTEM_PROMPT_TEMPLATE = """You are a specialist educational content creator producing self-study
booklets for {course_context}.

LANGUAGE: You MUST use UK English spellings throughout the entire booklet.
Examples: organise (not organize), colour (not color), centre (not center),
analyse (not analyze), labelled (not labeled), modelling (not modeling),
specialise (not specialize), defence (not defense), fibre (not fiber),
travelled (not traveled), minimise (not minimize), recognise (not recognize).
No section is exempt from this rule.

CRITICAL FORMATTING RULES — you MUST follow every one of these:

1. NEVER use double asterisks (**) anywhere in your output.  All emphasis
   must be conveyed through section headings or plain text — we apply bold
   formatting during docx conversion.  Using ** will break the pipeline.

2. STRUCTURE — every booklet must contain these sections IN ORDER:
   a. Title block (single-spaced, one line per field):
      Self-Study Booklet
      Subject: [Subject]
      Topic: [Topic]
      Specification Reference: [Ref]
      Lesson Number: [Number] (Year [Year])
      Lesson Title: [Title]
      Required Practical: [Yes/No — name if applicable]

   b. Section 1 — Holistic Recall Starter (20 questions, numbered 1-20)
   c. Section 2 — Key Vocabulary Table (Term | Definition)
   d. Section 3 — Knowledge Development (per chunk, see below)
   e. Section 4 — Drawing and Labelling (if applicable — instruct students to draw in their exercise book, do NOT leave blank space in the booklet)
   f. Section 5 — Calculations / Application Questions (numbered starting at 1)
   g. Section 6 — Summary / Key Takeaways (bullet points)
   h. Section 7 — Sentence Starters / Writing Frames (bullet points to help
      students structure written answers, e.g. "The function of the mitochondria is…")
   i. Section 8 — Mark Scheme (numbering restarts per section — see below)
   j. Section 9 — Self-Assessment / Progress Grid (RAG rating)
   k. Section 10 — Topics to Revisit (numbered 1, 2, 3 only)
   l. Section 11 — Targets for Next Lesson (numbered 1, 2, 3 only)
   m. Section 12 — Self-Assessment Actions (checkboxes)
   n. Section 13 — Document Info (version, date, page count)

3. KNOWLEDGE CHUNKS (Section 3) — create 3-5 chunks, each containing:
   - Introduction paragraph
   - Key Vocabulary table
   - Knowledge Content — BULLET POINTS ONLY (use - prefix), NEVER numbered
   - Worked Example (### heading) — BULLET POINTS ONLY, NEVER numbered.
     Adapt to the subject: source analysis for History, calculations for
     Maths/Science, case studies for Geography, close reading for English, etc.
     IMPORTANT: If the worked example references a source, quotation, data set,
     equation, diagram description, or any stimulus material, you MUST include
     the full text of that material in the booklet so students can use it
     independently. Never reference a source without reproducing it in full.
   - Misconception Box (### heading) — a common misconception students hold
     about this topic, with a clear explanation of the correct understanding
   - Knowledge Check Questions — numbered starting at 1 for EACH chunk
     (Chunk 2 questions start at 1, NOT continuing from Chunk 1)

   IMPORTANT — Worked Example and Misconception Box MUST appear as ### headings
   inside every Knowledge Chunk. Write them exactly as:
     ### Worked Example
     ### Misconception Box
   These headings trigger special coloured formatting in the output document.
   Never omit them and never write them as plain text labels.

4. UNIVERSAL NUMBERING RESTART RULE: Every distinct numbered section MUST
   start its own numbering at 1.  No numbered section should EVER continue
   numbering from a previous section.  This applies to:
   - Knowledge Check Questions (restart at 1 per chunk)
   - Application Questions (start at 1)
   - Section 9: Topics to Revisit (1, 2, 3)
   - Section 10: Targets for Next Lesson (1, 2, 3)
   The ONLY exception is the Holistic Recall Starter (1-20 continuous).

5. MARK SCHEME (Section 8) — CRITICAL NUMBERING RULES:
   a. Every mark scheme sub-section heading MUST include the words "Mark Scheme",
      e.g. "### Holistic Recall — Mark Scheme", "### Knowledge Chunk 1 — Mark Scheme",
      "### Application Questions — Mark Scheme".  This is mandatory.
   b. Numbering MUST exactly mirror the corresponding question section:
      - Holistic Recall — Mark Scheme:          entries 1–20 (matching questions 1–20)
      - Knowledge Chunk 1 — Mark Scheme:        entries 1, 2, 3 … (matching KC1 questions)
      - Knowledge Chunk 2 — Mark Scheme:        entries 1, 2, 3 … RESTART — NEVER continue from KC1
      - Application Questions — Mark Scheme:    entries 1, 2, 3 … RESTART
   c. CORRESPONDENCE: if Question 3 in Chunk 1 asks about X, then entry 3 in
      "Knowledge Chunk 1 — Mark Scheme" MUST answer X.  Check this for every entry.
   d. The number of mark scheme entries for each sub-section MUST equal the number
      of questions in that sub-section.  Count them before writing.

6. Use tables where appropriate (vocabulary, mark schemes, self-assessment).
7. Clearly mark Higher-Tier-only content with [HT ONLY] tags.
8. Target both Foundation and Higher tier.
9. Assume standard KS3 prior knowledge plus content from preceding lessons.

HEADING HIERARCHY — you MUST follow this exact markdown heading structure:
- # (H1) for the title line and every main section (Section 1, Section 2, etc.)
- ## (H2) for Knowledge Chunks within Section 3 (e.g. ## Knowledge Chunk 1: ...)
- ### (H3) for sub-sections within each chunk: ### Worked Example, ### Misconception Box
- ### (H3) also for Mark Scheme sub-sections (e.g. ### Knowledge Chunk 1 — Mark Scheme)
Do NOT shift headings down (e.g. using ## for main sections). Main sections MUST use #.

Output the complete booklet in well-structured markdown following these heading levels exactly."""


def get_system_prompt(course_config=None):
    """Build the system prompt, customised for the active course."""
    if course_config:
        context = course_config.get(
            "system_prompt_context", DEFAULT_SYSTEM_PROMPT_CONTEXT
        )
    else:
        context = DEFAULT_SYSTEM_PROMPT_CONTEXT
    return SYSTEM_PROMPT_TEMPLATE.format(course_context=context)


# Keep backwards-compatible reference for any existing imports
SYSTEM_PROMPT = SYSTEM_PROMPT_TEMPLATE.format(
    course_context=DEFAULT_SYSTEM_PROMPT_CONTEXT
)


# ---------------------------------------------------------------------------
# Claude API
# ---------------------------------------------------------------------------

def get_client():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError(
            "ANTHROPIC_API_KEY not set. Add it to .env file. "
            "Get a key from console.anthropic.com"
        )
    return anthropic.Anthropic(api_key=api_key)


def generate_booklet(lesson, prompt_text, model="claude-sonnet-4-5-20250929",
                     course_config=None):
    """
    Send a booklet generation request to Claude API.

    Uses prompt caching on the system prompt to save ~90% on input tokens
    across multiple booklet generations (the system prompt is identical
    for all booklets in a course).

    Model selection: Sonnet 4.5 is the optimal choice — same price as
    Sonnet 4 ($3/$15 per MTok) but higher quality output. Haiku would
    risk lower quality on complex structured content. Opus is overkill
    for templated generation.
    """
    from ai_client import create_message
    start = time.time()

    system_prompt = get_system_prompt(course_config)

    message = create_message(
        model=model,
        max_tokens=16000,
        system=[
            {
                "type": "text",
                "text": system_prompt,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[{"role": "user", "content": prompt_text}],
    )

    duration = time.time() - start
    content = ""
    for block in message.content:
        if block.type == "text":
            content += block.text

    # Extract cache performance info
    usage = {
        "input_tokens": message.usage.input_tokens,
        "output_tokens": message.usage.output_tokens,
    }
    # Include cache hit info if available
    if hasattr(message.usage, "cache_creation_input_tokens"):
        usage["cache_creation_input_tokens"] = message.usage.cache_creation_input_tokens
    if hasattr(message.usage, "cache_read_input_tokens"):
        usage["cache_read_input_tokens"] = message.usage.cache_read_input_tokens

    return {
        "content": content,
        "usage": usage,
        "model": message.model,
        "duration_s": round(duration, 1),
        "stop_reason": message.stop_reason,
    }


# ---------------------------------------------------------------------------
# Markdown sanitisation
# ---------------------------------------------------------------------------

def sanitize_markdown(content):
    """
    Post-process Claude's markdown to guarantee formatting compliance.
    Applied BEFORE docx conversion.

    Handles:
    - Stripping all ** markers
    - Forcing bullets in knowledge_content / worked_example sections
    - Restarting numbering in ALL numbered sections (universal rule)
    - UK English spelling corrections
    - Stripping any residual [DRAWING SPACE] markers
    """
    # --- UK English corrections (applied globally) ---
    content = _fix_uk_english(content)

    lines = content.split("\n")
    result = []
    in_section = None  # track which section we're in
    in_mark_scheme_block = False  # True once inside the overall mark scheme section
    q_counter = 0

    for line in lines:
        stripped = line.strip()

        # --- Detect current section from headings ---
        # Order matters: more specific checks (mark scheme) before broader ones
        # (knowledge chunk). "Knowledge Chunk 2 — Mark Scheme" must hit mark_scheme.
        heading_lower = stripped.lstrip("#").strip().lower()
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))

            # A new level-1 heading (main section) that isn't the mark scheme
            # exits the block. Level-2+ headings are sub-sections that may stay inside.
            if level == 1 and "mark scheme" not in heading_lower:
                in_mark_scheme_block = False

            if "mark scheme" in heading_lower:
                # Covers: "# Section 8 — Mark Scheme", "## KC1 — Mark Scheme",
                # "### Holistic Recall — Mark Scheme", etc.
                in_section = "mark_scheme"
                in_mark_scheme_block = True
                q_counter = 0
            elif in_mark_scheme_block:
                # Any sub-heading within the mark scheme block (e.g. "### Knowledge
                # Chunk 1" without "Mark Scheme" in title) — restart counter and
                # keep renumbering so the output always corresponds to the questions.
                in_section = "mark_scheme"
                q_counter = 0
            elif any(kw in heading_lower for kw in [
                "knowledge content", "key components", "key points",
                "knowledge development"
            ]):
                in_section = "knowledge_content"
            elif "worked example" in heading_lower:
                in_section = "worked_example"
            elif "knowledge check" in heading_lower:
                in_section = "knowledge_check"
                q_counter = 0
            elif any(kw in heading_lower for kw in [
                "application question", "calculation"
            ]):
                in_section = "application_questions"
                q_counter = 0
            elif any(kw in heading_lower for kw in [
                "topics to revisit", "topics for revisit"
            ]):
                in_section = "topics_to_revisit"
                q_counter = 0
            elif any(kw in heading_lower for kw in [
                "targets for next", "next lesson target"
            ]):
                in_section = "targets_next_lesson"
                q_counter = 0
            elif any(kw in heading_lower for kw in [
                "knowledge chunk", "chunk "
            ]):
                # New chunk heading — reset counter; sub-sections detected below
                q_counter = 0
                in_section = None
            else:
                in_section = None

        # --- Strip all ** markers ---
        line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)

        # --- Force bullets in knowledge_content and worked_example ---
        if in_section in ("knowledge_content", "worked_example"):
            m = re.match(r"^(\s*)\d+\.\s+(.+)", line)
            if m:
                line = f"{m.group(1)}- {m.group(2)}"

        # --- Universal numbering restart for all numbered sections ---
        if in_section in (
            "knowledge_check", "application_questions",
            "topics_to_revisit", "targets_next_lesson",
            "mark_scheme"
        ):
            m = re.match(r"^(\s*)\d+\.\s+(.+)", line)
            if m:
                q_counter += 1
                line = f"{m.group(1)}{q_counter}. {m.group(2)}"

        # --- Strip any [DRAWING SPACE] markers (diagrams removed) ---
        if not re.match(r"^\[DRAWING\s+SPACE", stripped, re.IGNORECASE):
            result.append(line)

    return "\n".join(result)


# Common US → UK English spelling substitutions for science booklets
_UK_SPELLING_FIXES = [
    # -ize → -ise
    (r"\borganize\b", "organise"), (r"\bOrganize\b", "Organise"),
    (r"\borganized\b", "organised"), (r"\bOrganized\b", "Organised"),
    (r"\borganizing\b", "organising"), (r"\bOrganizing\b", "Organising"),
    (r"\brecognize\b", "recognise"), (r"\bRecognize\b", "Recognise"),
    (r"\brecognized\b", "recognised"), (r"\bRecognized\b", "Recognised"),
    (r"\brecognizing\b", "recognising"), (r"\bRecognizing\b", "Recognising"),
    (r"\bminimize\b", "minimise"), (r"\bMinimize\b", "Minimise"),
    (r"\bminimized\b", "minimised"), (r"\bminimizing\b", "minimising"),
    (r"\bmaximize\b", "maximise"), (r"\bMaximize\b", "Maximise"),
    (r"\bmaximized\b", "maximised"), (r"\bmaximizing\b", "maximising"),
    (r"\bspecialize\b", "specialise"), (r"\bSpecialize\b", "Specialise"),
    (r"\bspecialized\b", "specialised"), (r"\bSpecialized\b", "Specialised"),
    (r"\bspecializing\b", "specialising"),
    (r"\butilize\b", "utilise"), (r"\bUtilize\b", "Utilise"),
    (r"\butilized\b", "utilised"), (r"\butilizing\b", "utilising"),
    (r"\banalyze\b", "analyse"), (r"\bAnalyze\b", "Analyse"),
    (r"\banalyzed\b", "analysed"), (r"\banalyzing\b", "analysing"),
    (r"\bsummarize\b", "summarise"), (r"\bSummarize\b", "Summarise"),
    (r"\bsummarized\b", "summarised"), (r"\bsummarizing\b", "summarising"),
    (r"\bmemorize\b", "memorise"), (r"\bMemorize\b", "Memorise"),
    (r"\bmemorized\b", "memorised"), (r"\bmemorizing\b", "memorising"),
    (r"\bvaporize\b", "vaporise"), (r"\bvaporization\b", "vaporisation"),
    (r"\bneutralize\b", "neutralise"), (r"\bneutralized\b", "neutralised"),
    (r"\bneutralizing\b", "neutralising"), (r"\bneutralization\b", "neutralisation"),
    (r"\boxidize\b", "oxidise"), (r"\boxidized\b", "oxidised"),
    (r"\boxidizing\b", "oxidising"),
    (r"\bcauterize\b", "cauterise"), (r"\bcauterized\b", "cauterised"),
    (r"\bionize\b", "ionise"), (r"\bionized\b", "ionised"),
    (r"\bionizing\b", "ionising"), (r"\bionization\b", "ionisation"),
    (r"\bpolymerize\b", "polymerise"), (r"\bpolymerized\b", "polymerised"),
    (r"\bpolymerization\b", "polymerisation"),
    (r"\bcatalyze\b", "catalyse"), (r"\bcatalyzed\b", "catalysed"),
    (r"\bcatalyzing\b", "catalysing"),
    (r"\bhydrolyze\b", "hydrolyse"), (r"\bhydrolyzed\b", "hydrolysed"),
    (r"\bhydrolyzing\b", "hydrolysing"), (r"\bhydrolysis\b", "hydrolysis"),
    (r"\bcustomize\b", "customise"), (r"\bcustomized\b", "customised"),
    # -or → -our
    (r"\bcolor\b", "colour"), (r"\bColor\b", "Colour"),
    (r"\bcolors\b", "colours"), (r"\bColors\b", "Colours"),
    (r"\bcolored\b", "coloured"),
    (r"\bfavor\b", "favour"), (r"\bFavor\b", "Favour"),
    (r"\bfavored\b", "favoured"), (r"\bfavoring\b", "favouring"),
    (r"\bfavorite\b", "favourite"), (r"\bFavorite\b", "Favourite"),
    (r"\bhonor\b", "honour"), (r"\bHonor\b", "Honour"),
    (r"\bhumor\b", "humour"), (r"\bHumor\b", "Humour"),
    (r"\bbehavior\b", "behaviour"), (r"\bBehavior\b", "Behaviour"),
    (r"\bbehaviors\b", "behaviours"),
    (r"\bneighbor\b", "neighbour"), (r"\bNeighbor\b", "Neighbour"),
    (r"\bneighbors\b", "neighbours"),
    (r"\blabor\b", "labour"), (r"\bLabor\b", "Labour"),
    (r"\bvapor\b", "vapour"), (r"\bVapor\b", "Vapour"),
    (r"\btumor\b", "tumour"), (r"\bTumor\b", "Tumour"),
    (r"\bodor\b", "odour"), (r"\bOdor\b", "Odour"),
    # -er → -re
    (r"\bcenter\b", "centre"), (r"\bCenter\b", "Centre"),
    (r"\bcenters\b", "centres"), (r"\bCenters\b", "Centres"),
    (r"\bcentered\b", "centred"),
    (r"\bfiber\b", "fibre"), (r"\bFiber\b", "Fibre"),
    (r"\bfibers\b", "fibres"),
    (r"\bliter\b", "litre"), (r"\bLiter\b", "Litre"),
    (r"\bliters\b", "litres"),
    (r"\bmeter\b", "metre"), (r"\bMeter\b", "Metre"),
    (r"\bmeters\b", "metres"),
    # -ed / -ing / -ling
    (r"\blabeled\b", "labelled"), (r"\bLabeled\b", "Labelled"),
    (r"\blabeling\b", "labelling"), (r"\bLabeling\b", "Labelling"),
    (r"\bmodeled\b", "modelled"), (r"\bModeled\b", "Modelled"),
    (r"\bmodeling\b", "modelling"), (r"\bModeling\b", "Modelling"),
    (r"\btraveled\b", "travelled"), (r"\btraveling\b", "travelling"),
    (r"\bcanceled\b", "cancelled"), (r"\bcanceling\b", "cancelling"),
    # -ence / -ense
    (r"\bdefense\b", "defence"), (r"\bDefense\b", "Defence"),
    (r"\boffense\b", "offence"), (r"\bOffense\b", "Offence"),
    (r"\blicense\b", "licence"), (r"\bLicense\b", "Licence"),
    (r"\bpractice\b(?=\s+(?:the|this|that|a|an))", "practise"),  # verb form only
    # Other common differences
    (r"\bgray\b", "grey"), (r"\bGray\b", "Grey"),
    (r"\bgrays\b", "greys"),
    (r"\bfetus\b", "foetus"), (r"\bFetus\b", "Foetus"),
    (r"\besophagus\b", "oesophagus"), (r"\bEsophagus\b", "Oesophagus"),
    (r"\bestrogen\b", "oestrogen"), (r"\bEstrogen\b", "Oestrogen"),
    (r"\bhemoglobin\b", "haemoglobin"), (r"\bHemoglobin\b", "Haemoglobin"),
    (r"\banemia\b", "anaemia"), (r"\bAnemia\b", "Anaemia"),
    (r"\bdiarrhea\b", "diarrhoea"), (r"\bDiarrhea\b", "Diarrhoea"),
    (r"\bpediatrician\b", "paediatrician"),
    (r"\bsulfur\b", "sulphur"), (r"\bSulfur\b", "Sulphur"),
    (r"\bsulfate\b", "sulphate"), (r"\bSulfate\b", "Sulphate"),
    (r"\bsulfide\b", "sulphide"), (r"\bSulfide\b", "Sulphide"),
    (r"\bsulfuric\b", "sulphuric"), (r"\bSulfuric\b", "Sulphuric"),
    (r"\baluminum\b", "aluminium"), (r"\bAluminum\b", "Aluminium"),
]


def _fix_uk_english(text):
    """Apply UK English spelling corrections to text."""
    for pattern, replacement in _UK_SPELLING_FIXES:
        text = re.sub(pattern, replacement, text)
    return text


# ---------------------------------------------------------------------------
# Markdown → Docx conversion  (Booklet Formatting Specification v1.0)
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    """Set border on a table cell.  kwargs: top, bottom, left, right each a dict."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, hex_color):
    """Set background shading on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def _set_cell_padding(cell, pt_val=6):
    """Set uniform cell padding in points."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(int(pt_val * 20)))  # twips
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_valign(cell, val="top"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    tcPr.append(vAlign)


def _set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _set_row_header(row):
    """Mark a table row as a repeating header row."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _para_border(paragraph, edge, sz, color):
    """Add a border to a paragraph edge ('top' or 'bottom')."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{edge}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(sz))
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


def _add_page_number(paragraph):
    """Insert a PAGE field into a paragraph (for footer)."""
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    run._r.append(fc1)

    run2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    run2._r.append(it)

    run3 = paragraph.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run3._r.append(fc2)


def _style_footer_runs(paragraph):
    for run in paragraph.runs:
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.color.rgb = COL_MUTED


def _add_section_divider(doc, page_break=False):
    """Horizontal rule before an H2 (or standalone ---).  Optionally forces a
    page break so the divider + heading land on the new page together."""
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # Tiny invisible run so paragraph height is minimal
    r = p.add_run()
    r.font.size = Pt(2)
    _para_border(p, "bottom", "8", HEX_BORDER_MD)  # 1 pt


def _add_formatted_text(paragraph, text):
    """Parse inline markdown formatting and add runs to paragraph."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    parts = re.split(r"(\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = FONT
            run.font.size = Pt(11)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part:
            run = paragraph.add_run(part)
            run.font.name = FONT
            run.font.size = Pt(11)


# ---- boxed element renderers ------------------------------------------------

def _render_misconception_box(doc, lines):
    """Red-bordered misconception box as a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.columns[0].width = Mm(TEXT_AREA_MM)
    cell = tbl.cell(0, 0)
    cell.text = ""
    ba = {"val": "single", "sz": "12", "color": HEX_MISC_BDR, "space": "0"}
    _set_cell_border(cell, top=ba, bottom=ba, left=ba, right=ba)
    _set_cell_shading(cell, HEX_MISC_BG)
    _set_cell_padding(cell, 10)

    first = True
    for raw in lines:
        txt = raw.strip()
        if not txt:
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.15

        upper = txt.upper()
        if upper.startswith("MISCONCEPTION:"):
            body = txt[len("MISCONCEPTION:"):].strip()
            r = p.add_run("MISCONCEPTION: ")
            r.bold = True; r.font.name = FONT; r.font.size = Pt(11)
            r.font.color.rgb = COL_MISC_TTL
            if body:
                r2 = p.add_run(body)
                r2.font.name = FONT; r2.font.size = Pt(11)
                r2.font.color.rgb = COL_PRIMARY
        elif upper.startswith("REALITY:"):
            body = txt[len("REALITY:"):].strip()
            r = p.add_run("REALITY: ")
            r.bold = True; r.font.name = FONT; r.font.size = Pt(11)
            r.font.color.rgb = COL_PRIMARY
            if body:
                r2 = p.add_run(body)
                r2.font.name = FONT; r2.font.size = Pt(11)
                r2.font.color.rgb = COL_PRIMARY
        else:
            _add_formatted_text(p, txt)

    if cell.paragraphs:
        cell.paragraphs[-1].paragraph_format.keep_with_next = False
    _set_row_cant_split(tbl.rows[0])


def _render_worked_example_box(doc, lines):
    """Blue-bordered worked-example box as a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.columns[0].width = Mm(TEXT_AREA_MM)
    cell = tbl.cell(0, 0)
    cell.text = ""
    ba = {"val": "single", "sz": "12", "color": HEX_WE_BDR, "space": "0"}
    _set_cell_border(cell, top=ba, bottom=ba, left=ba, right=ba)
    _set_cell_shading(cell, HEX_WE_BG)
    _set_cell_padding(cell, 10)

    first = True
    for raw in lines:
        txt = raw.strip()
        if not txt:
            continue

        # Bullets inside a worked example
        if txt.startswith("- ") or txt.startswith("* "):
            inner = txt[2:]
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.left_indent = Mm(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_together = True
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run("\u2022 ")
            r.font.name = FONT; r.font.size = Pt(11)
            _add_formatted_text(p, inner)
            continue

        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 1.15

        lower = txt.lower()
        if lower.startswith("question:"):
            body = txt[len("question:"):].strip()
            r = p.add_run("Question: ")
            r.bold = True; r.font.name = FONT; r.font.size = Pt(11)
            r.font.color.rgb = COL_WE_LABEL
            if body:
                r2 = p.add_run(body)
                r2.font.name = FONT; r2.font.size = Pt(11)
                r2.font.color.rgb = COL_PRIMARY
        elif lower.startswith("answer:"):
            body = txt[len("answer:"):].strip()
            r = p.add_run("Answer: ")
            r.bold = True; r.font.name = FONT; r.font.size = Pt(11)
            r.font.color.rgb = COL_WE_LABEL
            if body:
                r2 = p.add_run(body)
                r2.font.name = FONT; r2.font.size = Pt(11)
                r2.font.color.rgb = COL_PRIMARY
        else:
            _add_formatted_text(p, txt)

    if cell.paragraphs:
        cell.paragraphs[-1].paragraph_format.keep_with_next = False
    _set_row_cant_split(tbl.rows[0])


# ---- table renderer ----------------------------------------------------------

def _table_context(section):
    """Map current section to a table-rendering hint."""
    if section in ("key_vocabulary", "vocabulary"):
        return "vocabulary"
    if section == "self_assessment":
        return "self_assessment"
    return None


def _render_table(doc, table_rows, ctx=None):
    """Render a markdown table with specification formatting."""
    if not table_rows:
        return
    num_cols = max(len(r) for r in table_rows)
    num_rows = len(table_rows)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    tbl.style = "Table Grid"

    # Column widths
    if num_cols == 2 and ctx in ("vocabulary", None):
        tbl.columns[0].width = Mm(int(TEXT_AREA_MM * 0.30))
        tbl.columns[1].width = Mm(int(TEXT_AREA_MM * 0.70))
    elif num_cols == 4 and ctx == "self_assessment":
        tbl.columns[0].width = Mm(int(TEXT_AREA_MM * 0.55))
        for c in range(1, 4):
            tbl.columns[c].width = Mm(int(TEXT_AREA_MM * 0.15))
    else:
        cw = Mm(int(TEXT_AREA_MM / num_cols))
        for c in range(num_cols):
            tbl.columns[c].width = cw

    # Outer / inner borders via table-level XML
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    old_bdr = tblPr.find(qn("w:tblBorders"))
    if old_bdr is not None:
        tblPr.remove(old_bdr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), HEX_BORDER_MD); el.set(qn("w:space"), "0")
        tblBorders.append(el)
    for edge in ("insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), HEX_BORDER_LT); el.set(qn("w:space"), "0")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for r_idx, row_data in enumerate(table_rows):
        row = tbl.rows[r_idx]
        _set_row_cant_split(row)
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                continue
            cell = tbl.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            _set_cell_padding(cell, 6)
            _set_cell_valign(cell, "top")
            _add_formatted_text(p, cell_text)

            if r_idx == 0:
                _set_cell_shading(cell, HEX_TBL_HDR_BG)
                for run in p.runs:
                    run.bold = True
                    run.font.name = FONT
                    run.font.size = Pt(11)
                    run.font.color.rgb = COL_PRIMARY
                if ctx == "self_assessment" and c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                sz = Pt(10) if ctx == "self_assessment" else Pt(11)
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = sz
                    run.font.color.rgb = COL_PRIMARY
                if ctx == "self_assessment" and c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header row repeats
        if r_idx == 0:
            _set_row_header(row)
        # Keep small tables (<=6 body rows) on one page
        if num_rows <= 7 and r_idx < num_rows - 1:
            for p in row.cells[0].paragraphs:
                p.paragraph_format.keep_with_next = True


# ---- section detection -------------------------------------------------------

def _detect_section(heading_text):
    """Detect which section we're in from a heading."""
    h = heading_text.lower()
    if any(kw in h for kw in [
        "knowledge content", "key components", "key points",
        "knowledge development"
    ]):
        return "knowledge_content"
    if "worked example" in h:
        return "worked_example"
    if "misconception" in h:
        return "misconception_box"
    if "knowledge check" in h:
        return "knowledge_check"
    if any(kw in h for kw in ["application question", "calculation"]):
        return "application_questions"
    if any(kw in h for kw in ["topics to revisit", "topics for revisit"]):
        return "topics_to_revisit"
    if any(kw in h for kw in ["targets for next", "next lesson target"]):
        return "targets_next_lesson"
    if "mark scheme" in h:
        return "mark_scheme"
    if "key vocabulary" in h or ("vocabulary" in h and "table" not in h):
        return "key_vocabulary"
    if "self-assessment" in h and "action" not in h:
        return "self_assessment"
    if "progress grid" in h:
        return "self_assessment"
    return None


# ---- main converter ----------------------------------------------------------

def markdown_to_docx(md_path, lesson=None):
    """
    Convert a markdown booklet to .docx with full formatting
    per the Booklet Formatting Specification v1.0.
    """
    md_path = Path(md_path)
    content = md_path.read_text()

    doc = Document()

    # ── page setup ──────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        section.header_distance = Mm(12)
        section.footer_distance = Mm(12)
        section.gutter = Mm(0)

    # ── default font ────────────────────────────────────────────
    ns = doc.styles["Normal"]
    ns.font.name = FONT
    ns.font.size = Pt(11)
    ns.font.color.rgb = COL_PRIMARY
    ns.paragraph_format.line_spacing = 1.15
    ns.paragraph_format.space_after = Pt(6)
    ns.paragraph_format.widow_control = True

    # ── heading styles ──────────────────────────────────────────
    for level, (sz, col, bef, aft) in HEADING_SPEC.items():
        try:
            hs = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        hs.font.name = FONT
        hs.font.size = Pt(sz)
        hs.font.bold = True
        hs.font.color.rgb = col
        hs.paragraph_format.space_before = Pt(bef)
        hs.paragraph_format.space_after = Pt(aft)
        hs.paragraph_format.keep_with_next = True
        hs.paragraph_format.widow_control = True

    # ── header (different first page) ───────────────────────────
    sect = doc.sections[0]
    sect.different_first_page_header_footer = True

    # first-page header — empty (title page only)
    fh = sect.first_page_header
    fh.is_linked_to_previous = False

    # subsequent-pages header — lesson title + bottom rule
    hdr = sect.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_text = lesson.get("title", "") if lesson else ""
    hr = hp.add_run(title_text)
    hr.font.name = FONT; hr.font.size = Pt(9); hr.font.color.rgb = COL_MUTED
    _para_border(hp, "bottom", "4", HEX_BORDER_LT)

    # ── footer (page number on every page) ──────────────────────
    for ftr_obj in (sect.first_page_footer, sect.footer):
        ftr_obj.is_linked_to_previous = False
        fp = ftr_obj.paragraphs[0] if ftr_obj.paragraphs else ftr_obj.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(fp)
        _style_footer_runs(fp)
        _para_border(fp, "top", "4", HEX_BORDER_LT)

    # ── parse and build ─────────────────────────────────────────
    lines = content.split("\n")
    i = 0
    in_table = False
    table_rows = []
    in_title_block = False
    current_section = None
    box_type = None       # "misconception" | "worked_example"
    box_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # -- flush box on new heading --------------------------------
        if box_type and stripped.startswith("#"):
            if box_type == "misconception":
                _render_misconception_box(doc, box_lines)
            else:
                _render_worked_example_box(doc, box_lines)
            box_type = None
            box_lines = []
            # fall through to process heading

        # -- flush table when leaving table rows ---------------------
        if in_table and not (stripped.startswith("|") and "|" in stripped[1:]):
            in_table = False
            if table_rows:
                _render_table(doc, table_rows, _table_context(current_section))
            table_rows = []

        # -- headings ------------------------------------------------
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            hl = text.lower()
            current_section = _detect_section(text)

            # boxed-element headings — output H4 label then collect
            if "misconception" in hl and level >= 3:
                h = doc.add_heading(text, level=level)
                h.paragraph_format.keep_with_next = True
                for r in h.runs:
                    r.font.name = FONT
                box_type = "misconception"
                box_lines = []
                i += 1; continue
            if "worked example" in hl and level >= 3:
                h = doc.add_heading(text, level=level)
                h.paragraph_format.keep_with_next = True
                for r in h.runs:
                    r.font.name = FONT
                box_type = "worked_example"
                box_lines = []
                i += 1; continue

            # section divider + page break before every H2
            if level == 2:
                _add_section_divider(doc, page_break=True)

            h = doc.add_heading(text, level=min(level, 5))
            for r in h.runs:
                r.font.name = FONT

            # title-page detection
            if level == 1 and "self-study booklet" in hl:
                in_title_block = True
            i += 1; continue

        # -- collecting box content ----------------------------------
        if box_type:
            if stripped:
                box_lines.append(line)
            i += 1; continue

        # -- title block (metadata) ----------------------------------
        if in_title_block:
            if stripped == "":
                in_title_block = False
            else:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.15
                pf.keep_together = True
                pf.keep_with_next = True
                if ":" in stripped:
                    label, _, value = stripped.partition(":")
                    rl = p.add_run(label + ":")
                    rl.bold = True; rl.font.name = FONT; rl.font.size = Pt(10)
                    rl.font.color.rgb = COL_SECONDARY
                    rv = p.add_run(" " + value.strip())
                    rv.font.name = FONT; rv.font.size = Pt(10)
                    rv.font.color.rgb = COL_SECONDARY
                else:
                    _add_formatted_text(p, stripped)
                i += 1; continue

        # -- tables --------------------------------------------------
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells and not all(set(c) <= set("-: ") for c in cells):
                cells = [re.sub(r"\*\*([^*]+)\*\*", r"\1", c) for c in cells]
                table_rows.append(cells)
            i += 1; continue

        # -- horizontal rules ----------------------------------------
        if stripped in ("---", "***", "___"):
            _add_section_divider(doc, page_break=False)
            i += 1; continue

        # -- checkbox lists (Section 12) -----------------------------
        cb_match = None
        if stripped.startswith("\u2610 "):
            cb_match = stripped[2:]
        elif stripped.startswith("[ ] "):
            cb_match = stripped[4:]
        if cb_match is not None:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(8)
            p.paragraph_format.first_line_indent = Mm(-8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_together = True
            p.paragraph_format.widow_control = True
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run("\u2610 ")
            r.font.name = FONT; r.font.size = Pt(11)
            _add_formatted_text(p, cb_match)
            i += 1; continue

        # -- bullet points -------------------------------------------
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(8)
            p.paragraph_format.first_line_indent = Mm(-8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.widow_control = True
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run("\u2022 ")
            r.font.name = FONT; r.font.size = Pt(11)
            _add_formatted_text(p, text)
            i += 1; continue

        # -- numbered list -------------------------------------------
        num_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            if current_section in ("knowledge_content", "worked_example"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Mm(8)
                p.paragraph_format.first_line_indent = Mm(-8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.widow_control = True
                p.paragraph_format.line_spacing = 1.15
                r = p.add_run("\u2022 ")
                r.font.name = FONT; r.font.size = Pt(11)
                _add_formatted_text(p, text)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Mm(8)
                p.paragraph_format.first_line_indent = Mm(-8)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_together = True
                p.paragraph_format.widow_control = True
                p.paragraph_format.line_spacing = 1.15
                _add_formatted_text(p, f"{num}. {text}")
            i += 1; continue

        # -- sub-part list (a), b), …) --------------------------------
        sub_match = re.match(r"^([a-z])\)\s+(.+)", stripped)
        if sub_match:
            letter = sub_match.group(1)
            text = sub_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(16)
            p.paragraph_format.first_line_indent = Mm(-8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.widow_control = True
            p.paragraph_format.line_spacing = 1.15
            _add_formatted_text(p, f"{letter}) {text}")
            i += 1; continue

        # -- empty line → skip (spacing from styles) -----------------
        if not stripped:
            i += 1; continue

        # -- regular paragraph ---------------------------------------
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.widow_control = True
        p.paragraph_format.line_spacing = 1.15
        _add_formatted_text(p, stripped)
        i += 1

    # ── flush remaining content ─────────────────────────────────
    if box_type:
        if box_type == "misconception":
            _render_misconception_box(doc, box_lines)
        else:
            _render_worked_example_box(doc, box_lines)
    if in_table and table_rows:
        _render_table(doc, table_rows, _table_context(current_section))

    # ── save ────────────────────────────────────────────────────
    docx_path = md_path.with_suffix(".docx")
    doc.save(str(docx_path))
    return str(docx_path)


# ---------------------------------------------------------------------------
# File naming & existence checks
# ---------------------------------------------------------------------------

def _build_filename(lesson, ext=".md"):
    """Build the standardised filename, including RP code if applicable."""
    num = lesson["lesson_number"]
    title = lesson["title"] or "Untitled"
    rp = lesson.get("required_practical") or ""

    # Only prepend RP code if title doesn't already start with it
    if rp and rp.lower() not in ("none", "n/a", "") and not title.lower().startswith(rp.lower()):
        fname = f"L{num:03d} - {rp} - {title}{ext}"
    else:
        fname = f"L{num:03d} - {title}{ext}"

    return fname.replace("/", "-").replace(":", " -")


def _build_output_dir(lesson):
    """Build and return the output directory for a lesson."""
    topic_folder = lesson.get("output_folder", "").rstrip("/")
    subject = lesson["subject"] or "Unknown"
    out_dir = OUTPUT_DIR / topic_folder if topic_folder else OUTPUT_DIR / subject
    return out_dir


def check_existing_booklet(lesson):
    """Check if a booklet already exists for this lesson.

    Checks the current output path first, then recursively searches older
    path layouts for backward compatibility.
    """
    out_dir = _build_output_dir(lesson)
    docx_name = _build_filename(lesson, ".docx")
    pdf_name = _build_filename(lesson, ".pdf")

    docx_path = out_dir / docx_name
    pdf_path = out_dir / pdf_name

    # Also check old naming convention (without RP code)
    old_docx_name = f"L{lesson['lesson_number']:03d} - {lesson['title']}.docx".replace("/", "-").replace(":", " -")
    old_docx = out_dir / old_docx_name

    if docx_path.exists():
        found_path = docx_path
    elif old_docx.exists():
        found_path = old_docx
    else:
        # Search the entire output directory for the filename.
        # This handles all legacy path formats (course_id-based, pre-multi-course, etc.)
        found_path = None
        for candidate in OUTPUT_DIR.rglob(docx_name):
            found_path = candidate
            break
        if not found_path:
            for candidate in OUTPUT_DIR.rglob(old_docx_name):
                found_path = candidate
                break

    if found_path is None:
        found_path = docx_path  # default (doesn't exist)

    return {
        "exists": found_path.exists(),
        "docx_path": str(found_path),
        "pdf_path": str(pdf_path),
    }


def delete_lesson_files_from_disk(lesson):
    """Delete all booklet files (.md, .docx, .pdf) for a lesson from disk.

    Uses check_existing_booklet() to find files regardless of path layout.
    Returns dict with 'deleted' (list of paths removed) and 'not_found' count.
    """
    result = {"deleted": [], "not_found": 0}

    existing = check_existing_booklet(lesson)
    if not existing["exists"]:
        return result

    docx_path = Path(existing["docx_path"])
    base_dir = docx_path.parent
    stem = docx_path.stem

    for ext in [".md", ".docx", ".pdf"]:
        target = base_dir / f"{stem}{ext}"
        if target.exists():
            target.unlink()
            result["deleted"].append(str(target))
        else:
            result["not_found"] += 1

    # Clean up empty parent directories up to OUTPUT_DIR
    try:
        parent = base_dir
        while parent != OUTPUT_DIR and parent.exists():
            if not any(p for p in parent.iterdir() if p.name != ".DS_Store"):
                # Remove .DS_Store too if it's the only thing left
                ds = parent / ".DS_Store"
                if ds.exists():
                    ds.unlink()
                parent.rmdir()
                parent = parent.parent
            else:
                break
    except OSError:
        pass

    return result


# ---------------------------------------------------------------------------
# Save markdown
# ---------------------------------------------------------------------------

def save_booklet_markdown(lesson, content):
    """Save generated content as markdown file in output directory."""
    out_dir = _build_output_dir(lesson)
    out_dir.mkdir(parents=True, exist_ok=True)

    fname = _build_filename(lesson, ".md")
    fpath = out_dir / fname
    fpath.write_text(content)
    return str(fpath)


# ---------------------------------------------------------------------------
# PDF conversion
# ---------------------------------------------------------------------------

def convert_to_pdf(docx_path):
    """
    Convert a .docx to .pdf using LibreOffice.

    Falls back to docx2pdf (which needs MS Word) if LibreOffice isn't available.
    Returns pdf path or None on failure.
    """
    import shutil
    import subprocess

    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")

    # Strategy 1: LibreOffice (free, works on Mac/Linux/Windows)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(docx_path.parent),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if pdf_path.exists():
                logger.info(f"PDF created via LibreOffice: {pdf_path}")
                return str(pdf_path)
            else:
                logger.warning(f"LibreOffice conversion produced no file. stderr: {result.stderr}")
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed: {e}")

    # Strategy 2: docx2pdf (needs MS Word on Mac)
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return str(pdf_path)
    except Exception as e:
        logger.warning(f"docx2pdf conversion failed: {e}")

    logger.warning("PDF conversion failed — install LibreOffice or MS Word")
    return None


# ---------------------------------------------------------------------------
# Full pipeline
# ---------------------------------------------------------------------------

def generate_and_save(lesson, prompt_text, model="claude-sonnet-4-5-20250929",
                      replace=False, course_config=None):
    """
    Full pipeline: generate via API → sanitise → docx → pdf.

    Args:
        lesson: lesson dict from parser
        prompt_text: the master prompt
        model: Claude model to use
        replace: if True, overwrite existing files
        course_config: optional course config dict

    Returns dict with paths and metadata.
    """
    # Check for existing files
    existing = check_existing_booklet(lesson)
    if existing["exists"] and not replace:
        raise FileExistsError(
            f"Booklet already exists at {existing['docx_path']}. "
            "Set replace=True to overwrite."
        )

    # Generate via Claude API
    result = generate_booklet(
        lesson, prompt_text, model=model, course_config=course_config
    )

    # Sanitise the markdown
    clean_content = sanitize_markdown(result["content"])

    # Save markdown
    md_path = save_booklet_markdown(lesson, clean_content)

    # Convert to docx
    docx_path = markdown_to_docx(md_path, lesson=lesson)

    # Convert to PDF
    pdf_path = convert_to_pdf(docx_path)

    return {
        "md_path": md_path,
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "usage": result["usage"],
        "model": result["model"],
        "duration_s": result["duration_s"],
        "stop_reason": result["stop_reason"],
    }
